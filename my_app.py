from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

#profile piture
document.add_picture(
    'CV.jpg', width=Inches(2.0)
)

#name phone number and emails details
name =  input('What is your name? ')
speak('Hello ' + name + 'I hope you are doing well')
speak('What is your phone number? ')
phone_number = input('What is your phone number? ')
speak('What is your email? ')
email = input('What is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

#about me
document.add_heading('About me')
speak('Tell me about yourself  ')
about_me = input('Tell me about yourself ')
document.add_paragraph(about_me)

# Otra forma de hacerlo es poniendo el input directamente como parámetro:
# document.add_paragraph(
#       input('Tell about yourself? ')
# )

# working experience
document.add_heading('Work Experience')
p = document.add_paragraph()

speak('Tell me about your work experience:  ')
company = input('Enter company ')
from_date = input ('From Date ')
to_date = input ('To Date ')

p.add_run(company + ' '). bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

speak('Describe your experience at ' + company )
experience_details = input(
    'Describe your experience at ' + company + ' ')
p.add_run(experience_details)

# more experiences
while True:
    speak('Do you hace more experiences? Yes or No ')
    has_more_experiences = input(
        'Do you hace more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        speak('Tell me about your work experience:  ')
        company = input('Enter company ')
        from_date = input ('From Date ')
        to_date = input ('To Date ')

        p.add_run(company + ' '). bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True
       
        speak('Describe your experience at ' + company )
        experience_details = input(
            'Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else: 
        break

# skills
document.add_heading('Skills')
speak('Enter skill' )
skill = input('Enter skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'


while True:
    speak('Do you hace more skills? Yes or No ')
    has_more_skills = input(
        'Do you hace more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        speak('Enter skill' )
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'

    else: 
        break


# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Amigoscode and Intuit QuickBooks course project'

document.save('cv.docx')